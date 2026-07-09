# main.py
import fitz
import re
import os
import subprocess
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
import platform

print("=" * 70)
print("AI APPLICATION GENERATOR")
print("=" * 70)
print("\nThis tool generates professional Cover Letters and Cold Emails")
print("based on your Resume and Job Description.\n")

# Load environment variables
load_dotenv()

# Get API key
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

if not GROQ_API_KEY:
    print("❌ Error: GROQ_API_KEY not found in .env file")
    print("Please make sure your .env file contains: GROQ_API_KEY=your_api_key_here")
    exit()

print("✅ API Key loaded successfully!")

# Initialize Groq client
try:
    client = Groq(api_key=GROQ_API_KEY)
    print("✅ Groq client initialized successfully!\n")
except Exception as e:
    print(f"❌ Failed to initialize Groq client: {e}")
    exit()

def get_downloads_folder():
    """Get the user's Downloads folder path"""
    if platform.system() == 'Windows':
        return os.path.join(os.environ['USERPROFILE'], 'Downloads')
    elif platform.system() == 'Darwin':  # macOS
        return os.path.join(os.path.expanduser('~'), 'Downloads')
    else:  # Linux
        return os.path.join(os.path.expanduser('~'), 'Downloads')

def open_file(file_path):
    """Open file with default application"""
    try:
        if platform.system() == 'Windows':
            os.startfile(file_path)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.run(['open', file_path])
        else:  # Linux
            subprocess.run(['xdg-open', file_path])
        return True
    except Exception as e:
        print(f"⚠️ Could not open file: {e}")
        return False

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print("Error Reading PDF:", e)
        return None

def extract_jd_details(jd_text):
    """Extract company, role, and recruiter from job description"""
    company = None
    role = None
    recruiter = None

    company_match = re.search(
        r"Company\s*:\s*(.*)",
        jd_text,
        re.IGNORECASE
    )
    if company_match:
        company = company_match.group(1).strip()

    role_match = re.search(
        r"Job Title\s*:\s*(.*)",
        jd_text,
        re.IGNORECASE
    )
    if role_match:
        role = role_match.group(1).strip()

    recruiter_match = re.search(
        r"(Hiring Manager|HR|Recruitment Team|Hiring Team)",
        jd_text,
        re.IGNORECASE
    )
    if recruiter_match:
        recruiter = recruiter_match.group(1)

    return company, role, recruiter

def get_case_insensitive_input(prompt, valid_options=None):
    """
    Get user input and make it case-insensitive
    valid_options: list of valid choices (e.g., ['1', '2', '3'])
    """
    while True:
        user_input = input(prompt).strip()
        
        if valid_options is None:
            return user_input
            
        # Check if input matches any valid option (case-insensitive)
        if user_input.lower() in [opt.lower() for opt in valid_options]:
            return user_input.lower()
        else:
            print(f"❌ Invalid input. Please enter one of: {', '.join(valid_options)}")

def fix_contact_details(text):
    """Fix contact details to ensure each is on its own line"""
    # Fix Email, Phone, LinkedIn formatting
    # Pattern: "Name Email: email Phone: phone LinkedIn: linkedin"
    # Should be: "Name\nEmail: email\nPhone: phone\nLinkedIn: linkedin"
    
    # Find the closing section
    lines = text.split('\n')
    new_lines = []
    contact_found = False
    contact_line = ""
    
    for line in lines:
        if "Yours sincerely" in line or "Yours sincerely," in line:
            new_lines.append(line)
            contact_found = True
            continue
        
        if contact_found:
            # Check if this line contains contact details
            if "Email:" in line or "Phone:" in line or "LinkedIn:" in line or "@" in line or "+" in line:
                contact_line += " " + line.strip()
            else:
                new_lines.append(line)
        else:
            new_lines.append(line)
    
    # Process the contact line
    if contact_line:
        # Clean up the contact line
        contact_line = contact_line.strip()
        
        # Extract Email
        email_match = re.search(r'Email:\s*([^\s]+)', contact_line, re.IGNORECASE)
        email = email_match.group(0) if email_match else ""
        
        # Extract Phone
        phone_match = re.search(r'Phone:\s*([^\s]+(?:\s*[^\s]+)*)', contact_line, re.IGNORECASE)
        phone = phone_match.group(0) if phone_match else ""
        
        # Extract LinkedIn
        linkedin_match = re.search(r'LinkedIn:\s*([^\s]+)', contact_line, re.IGNORECASE)
        linkedin = linkedin_match.group(0) if linkedin_match else ""
        
        # If no labels found, try to extract from raw text
        if not email and '@' in contact_line:
            email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', contact_line)
            if email_match:
                email = f"Email: {email_match.group(1)}"
        
        if not phone and '+' in contact_line:
            phone_match = re.search(r'(\+\d{1,3}\s*\d{10})', contact_line)
            if phone_match:
                phone = f"Phone: {phone_match.group(1)}"
        
        if not linkedin and 'linkedin.com' in contact_line:
            linkedin_match = re.search(r'(linkedin\.com/[^\s]+)', contact_line)
            if linkedin_match:
                linkedin = f"LinkedIn: {linkedin_match.group(1)}"
        
        # Build the contact section
        contact_section = []
        if email:
            contact_section.append(email)
        if phone:
            contact_section.append(phone)
        if linkedin:
            contact_section.append(linkedin)
        
        # Replace the contact line with properly formatted lines
        if contact_section:
            # Find where the contact line was and replace it
            new_lines = [line for line in new_lines if not (('Email:' in line or 'Phone:' in line or 'LinkedIn:' in line or '@' in line or '+' in line) and not 'Yours sincerely' in line)]
            # Add the contact section
            for contact in contact_section:
                new_lines.append(contact)
    
    # Fix "www.linkedin.com" to "LinkedIn: www.linkedin.com"
    final_text = '\n'.join(new_lines)
    final_text = re.sub(r'www\.linkedin\.com', 'LinkedIn: www.linkedin.com', final_text, flags=re.IGNORECASE)
    final_text = re.sub(r'linkedin\.com', 'LinkedIn: linkedin.com', final_text, flags=re.IGNORECASE)
    
    return final_text

def format_cover_letter(text):
    """Format Cover Letter with proper professional structure"""
    # Remove any existing titles or headers
    text = re.sub(r'COVER LETTER', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Cover Letter', '', text, flags=re.IGNORECASE)
    
    # Fix "Dear hr" to "Dear Hiring Manager"
    text = re.sub(r'Dear\s+hr[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Dear\s+Hiring\s+Manager[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    
    # Fix closing
    text = re.sub(r'Yours sincerely,?\s*([A-Z][a-z]+ [A-Z][a-z]+)', r'Yours sincerely,\n\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sYours sincerely', 'Yours sincerely', text, flags=re.IGNORECASE)
    
    # Fix contact details
    text = fix_contact_details(text)
    
    # Remove any separator text
    text = re.sub(r'---\s*SEPARATOR\s*---', '', text, flags=re.IGNORECASE)
    text = re.sub(r'SEPARATOR', '', text, flags=re.IGNORECASE)
    text = re.sub(r'-{3,}', '', text)
    
    # Add proper spacing after periods for paragraphs
    text = re.sub(r'\. ([A-Z])', r'. \n\n\1', text)
    
    # Ensure proper spacing between sections
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove duplicate www. in LinkedIn
    text = re.sub(r'www\.\s*LinkedIn:', 'LinkedIn:', text, flags=re.IGNORECASE)
    
    # Clean up extra spaces
    text = re.sub(r' +', ' ', text)
    
    return text.strip()

def format_cold_email(text):
    """Format Cold Email with proper professional structure"""
    # Remove any existing titles or headers
    text = re.sub(r'COLD EMAIL', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Cold Email', '', text, flags=re.IGNORECASE)
    
    # Ensure Subject: is on its own line
    text = re.sub(r'Subject:\s*', 'Subject: ', text, flags=re.IGNORECASE)
    if not text.startswith('Subject:'):
        text = 'Subject: Application for Position\n\n' + text
    
    # Fix "Dear hr" to "Dear Hiring Manager"
    text = re.sub(r'Dear\s+hr[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Dear\s+Hiring\s+Manager[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    
    # Fix closing
    text = re.sub(r'Yours sincerely,?\s*([A-Z][a-z]+ [A-Z][a-z]+)', r'Yours sincerely,\n\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sYours sincerely', 'Yours sincerely', text, flags=re.IGNORECASE)
    
    # Fix contact details
    text = fix_contact_details(text)
    
    # Remove any separator text
    text = re.sub(r'---\s*SEPARATOR\s*---', '', text, flags=re.IGNORECASE)
    text = re.sub(r'SEPARATOR', '', text, flags=re.IGNORECASE)
    text = re.sub(r'-{3,}', '', text)
    
    # Ensure proper spacing between sections
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove duplicate www. in LinkedIn
    text = re.sub(r'www\.\s*LinkedIn:', 'LinkedIn:', text, flags=re.IGNORECASE)
    
    # Clean up extra spaces
    text = re.sub(r' +', ' ', text)
    
    return text.strip()

def save_as_docx(content, filename, content_type="Cover Letter"):
    """Save content as DOCX file with professional formatting"""
    try:
        doc = Document()
        
        # Add professional margins
        section = doc.sections[0]
        section.top_margin = 72  # 1 inch
        section.bottom_margin = 72
        section.left_margin = 72
        section.right_margin = 72
        
        # Format based on content type
        if "Cover Letter" in content_type:
            content = format_cover_letter(content)
        else:
            content = format_cold_email(content)
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if it's a date (left aligned)
                if re.match(r'^[A-Za-z]+\s+\d{1,2},?\s+\d{4}$', para.strip()):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Check if it's "Dear" line (left aligned)
                elif para.strip().startswith('Dear'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Check if it's "Subject:" line (left aligned)
                elif para.strip().startswith('Subject:'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Check if it's "Yours sincerely," (left aligned)
                elif para.strip().startswith('Yours sincerely'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Check if it's contact details (left aligned)
                elif any(keyword in para.strip() for keyword in ['Email:', 'Phone:', 'LinkedIn:']):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Regular paragraph with indentation
                else:
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = 36  # 0.5 inch indent
                    p.paragraph_format.space_after = 12
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Save document
        doc.save(filename)
        print(f"✅ {content_type} saved as: {filename}")
        return True
    except Exception as e:
        print(f"❌ Error saving {content_type} as DOCX: {e}")
        return False

def save_as_pdf(content, filename, content_type="Cover Letter"):
    """Save content as PDF file with professional formatting"""
    try:
        doc = SimpleDocTemplate(
            filename, 
            pagesize=letter,
            leftMargin=72,
            rightMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        styles = getSampleStyleSheet()
        
        # Create custom styles
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            leading=14,
            firstLineIndent=36  # 0.5 inch indent for paragraphs
        )
        
        left_style = ParagraphStyle(
            'LeftAlign',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            leading=14
        )
        
        # Format based on content type
        if "Cover Letter" in content_type:
            content = format_cover_letter(content)
        else:
            content = format_cold_email(content)
        
        # Build story
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check if it's a date (left aligned, no indent)
                if re.match(r'^[A-Za-z]+\s+\d{1,2},?\s+\d{4}$', para.strip()):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                # Check if it's "Dear" line (left aligned)
                elif para.strip().startswith('Dear'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                # Check if it's "Subject:" line (left aligned)
                elif para.strip().startswith('Subject:'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                # Check if it's "Yours sincerely," (left aligned)
                elif para.strip().startswith('Yours sincerely'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                # Check if it's contact details (left aligned)
                elif any(keyword in para.strip() for keyword in ['Email:', 'Phone:', 'LinkedIn:']):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.05*inch))
                else:
                    # Regular paragraph with indent
                    story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 0.05*inch))
        
        # Build PDF
        doc.build(story)
        print(f"✅ {content_type} saved as: {filename}")
        return True
    except Exception as e:
        print(f"❌ Error saving {content_type} as PDF: {e}")
        return False

def save_cover_letter(content, output_dir):
    """Save cover letter in both DOCX and PDF formats"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    docx_filename = os.path.join(output_dir, f"cover_letter_{timestamp}.docx")
    pdf_filename = os.path.join(output_dir, f"cover_letter_{timestamp}.pdf")
    
    docx_success = save_as_docx(content, docx_filename, "Cover Letter")
    pdf_success = save_as_pdf(content, pdf_filename, "Cover Letter")
    
    if docx_success and pdf_success:
        return {
            'docx': docx_filename,
            'pdf': pdf_filename
        }
    return None

def save_cold_email(content, output_dir):
    """Save cold email in both DOCX and PDF formats"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    docx_filename = os.path.join(output_dir, f"cold_email_{timestamp}.docx")
    pdf_filename = os.path.join(output_dir, f"cold_email_{timestamp}.pdf")
    
    docx_success = save_as_docx(content, docx_filename, "Cold Email")
    pdf_success = save_as_pdf(content, pdf_filename, "Cold Email")
    
    if docx_success and pdf_success:
        return {
            'docx': docx_filename,
            'pdf': pdf_filename
        }
    return None

def extract_cover_letter_and_email(content):
    """Extract Cover Letter and Cold Email from combined content"""
    cover_letter = None
    cold_email = None
    
    # Try to split by "--- SEPARATOR ---"
    if "SEPARATOR" in content:
        parts = re.split(r'---?\s*SEPARATOR\s*---?', content, flags=re.IGNORECASE)
        if len(parts) >= 2:
            cover_letter = parts[0].strip()
            cold_email = parts[1].strip()
            return cover_letter, cold_email
    
    # If no separator, try to split by "Subject:"
    if "Subject:" in content:
        parts = content.split("Subject:", 1)
        if len(parts) >= 2:
            cover_letter = parts[0].strip()
            cold_email = "Subject:" + parts[1].strip()
            return cover_letter, cold_email
    
    # Try to split by multiple "Dear" occurrences
    dear_indices = []
    lines = content.split('\n')
    for i, line in enumerate(lines):
        if "Dear" in line:
            dear_indices.append(i)
    
    if len(dear_indices) >= 2:
        cover_letter = '\n'.join(lines[:dear_indices[1]]).strip()
        cold_email = '\n'.join(lines[dear_indices[1]:]).strip()
        return cover_letter, cold_email
    
    # If all else fails, return the content as cover letter
    return content, None

def save_content(content, choice):
    """Save generated content to files based on choice"""
    # Create output directory in Downloads folder
    downloads_folder = get_downloads_folder()
    output_dir = os.path.join(downloads_folder, "AI_Generated_Documents")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📁 Created folder: {output_dir}")
    
    saved_files = {}
    
    if choice == "1":
        # Save only Cover Letter
        print("\n📄 Saving Cover Letter...")
        result = save_cover_letter(content, output_dir)
        if result:
            saved_files['cover_letter'] = result
        else:
            print("❌ Failed to save Cover Letter")
    
    elif choice == "2":
        # Save only Cold Email
        print("\n📄 Saving Cold Email...")
        result = save_cold_email(content, output_dir)
        if result:
            saved_files['cold_email'] = result
        else:
            print("❌ Failed to save Cold Email")
    
    elif choice == "3":
        # Save both Cover Letter and Cold Email separately
        print("\n📄 Extracting Cover Letter and Cold Email...")
        
        cover_letter, cold_email = extract_cover_letter_and_email(content)
        
        # Save Cover Letter
        if cover_letter:
            print("\n📄 Saving Cover Letter...")
            result = save_cover_letter(cover_letter, output_dir)
            if result:
                saved_files['cover_letter'] = result
        else:
            print("⚠️  Could not extract Cover Letter")
        
        # Save Cold Email
        if cold_email:
            print("\n📄 Saving Cold Email...")
            result = save_cold_email(cold_email, output_dir)
            if result:
                saved_files['cold_email'] = result
        else:
            print("⚠️  Could not extract Cold Email")
        
        # If both failed, save the entire content as both
        if not saved_files:
            print("\n⚠️  Could not separate content. Saving as both documents...")
            result = save_cover_letter(content, output_dir)
            if result:
                saved_files['cover_letter'] = result
            result = save_cold_email(content, output_dir)
            if result:
                saved_files['cold_email'] = result
    
    return saved_files, output_dir

def open_pdf_files(saved_files):
    """Open all PDF files that were saved"""
    opened_files = []
    
    for doc_type, files in saved_files.items():
        if 'pdf' in files:
            pdf_path = files['pdf']
            print(f"\n📂 Opening {doc_type.replace('_', ' ').upper()} PDF...")
            if open_file(pdf_path):
                opened_files.append(pdf_path)
    
    return opened_files

def main():
    print("\n" + "=" * 70)
    print("STEP 1: Provide your Resume")
    print("=" * 70)
    
    # Get resume path
    resume_path = input("\n📄 Enter Resume PDF Path: ")

    print("\n" + "=" * 70)
    print("STEP 2: Paste Job Description")
    print("=" * 70)
    print("\nPaste the complete Job Description below")
    print("Type 'DONE' in a new line when finished.\n")

    jd_lines = []
    while True:
        line = input()
        if line.strip().upper() == "DONE":
            break
        jd_lines.append(line)

    job_description = "\n".join(jd_lines)
    print("\n✅ Job Description Captured Successfully!\n")

    print("=" * 70)
    print("STEP 3: Extracting Details")
    print("=" * 70)
    
    # Extract details from JD
    company_name, job_role, recruiter_name = extract_jd_details(job_description)

    # Get missing details
    if not company_name:
        company_name = get_case_insensitive_input(
            "\n🏢 Company Name not found.\nEnter Company Name: "
        )

    if not job_role:
        job_role = get_case_insensitive_input(
            "\n💼 Job Role not found.\nEnter Job Role: "
        )

    if not recruiter_name:
        recruiter_name = get_case_insensitive_input(
            "\n👤 Recruiter Name not found.\n"
            "Enter Hiring Manager / HR / Hiring Team: "
        )

    print("\n" + "=" * 70)
    print("STEP 4: Select Document Type")
    print("=" * 70)
    print("\nPlease select what you want to generate:")
    print("  1. Cover Letter Only (Detailed, 4-5 paragraphs)")
    print("  2. Cold Email Only (Short, 80-100 words)")
    print("  3. Both (Cover Letter + Cold Email)")
    
    choice = get_case_insensitive_input(
        "\n📝 Enter your choice (1/2/3): ",
        valid_options=['1', '2', '3']
    )

    print("\n" + "=" * 70)
    print("STEP 5: Choose Tone")
    print("=" * 70)
    print("\nSelect the tone for your documents:")
    print("  - Professional (Formal, business-like)")
    print("  - Confident (Assertive, strong)")
    print("  - Enthusiastic (Energetic, passionate)")
    print("  - Fresher-Friendly (Entry-level, eager)")
    print("  - Experienced (Seasoned, expert)")
    
    valid_tones = ['professional', 'confident', 'enthusiastic', 'fresher-friendly', 'experienced']
    tone = get_case_insensitive_input(
        "\n🎯 Enter tone (Professional/Confident/Enthusiastic/Fresher-Friendly/Experienced): ",
        valid_options=valid_tones
    )

    # Extract resume text
    print("\n📖 Extracting text from resume...")
    resume_text = extract_text_from_pdf(resume_path)
    if resume_text is None:
        print("❌ Failed to read Resume PDF.")
        return
    print("✅ Resume extracted successfully!\n")

    print("=" * 70)
    print("STEP 6: Generating Content")
    print("=" * 70)
    print("\n⏳ This may take a few moments...\n")

    # Professional Cover Letter Prompt
    cover_letter_prompt = """
Generate a PROFESSIONAL COVER LETTER with proper business letter format.

EXACT STRUCTURE (Follow this exactly with proper spacing):

[Current Date]

[Recipient Full Name]
[Recipient Job Title]
[Company Name]
[Company Street Address]
[City, State, ZIP Code]

Dear Hiring Manager,

[Opening Paragraph - 2-3 sentences]
- Introduce yourself
- State the position you're applying for
- Express enthusiasm and interest

[Body Paragraph 1 - 3-4 sentences]
- Highlight your relevant skills and qualifications
- Connect your experience to the job requirements
- Use keywords from the job description

[Body Paragraph 2 - 3-4 sentences]
- Mention specific projects, achievements, or accomplishments
- Provide measurable results if possible
- Show how you've added value in previous roles

[Body Paragraph 3 - 3-4 sentences]
- Explain why you're interested in this company
- Research the company and mention something specific
- Show how you can contribute to their goals

[Closing Paragraph - 2-3 sentences]
- Thank the hiring manager for their time
- Express interest in an interview
- Mention attached resume and additional documents

Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.

Yours sincerely,

[Your Full Name]
Email: [Your Email]
Phone: [Your Phone Number]
LinkedIn: [Your LinkedIn URL]

CRITICAL FORMATTING RULES:
1. Each section on its own line with blank lines between
2. "Yours sincerely," on its own line with a blank line after
3. Name on its own line after "Yours sincerely,"
4. Email on its own line
5. Phone on its own line
6. LinkedIn on its own line
7. Use "Dear Hiring Manager," (not "Dear hr")
8. Professional left alignment
9. First line of each paragraph indented
10. 4-5 paragraphs in the body
"""

    # Professional Cold Email Prompt
    cold_email_prompt = """
Generate a PROFESSIONAL COLD EMAIL with proper business email format.

EXACT STRUCTURE (Follow this exactly with proper spacing):

Subject: [Clear, concise, and professional subject line]

Dear Hiring Manager,

[Opening - 1-2 sentences]
- Introduce yourself briefly
- State the position you're applying for
- Express interest

[Body - 2-3 sentences]
- Highlight your most relevant key skills
- Mention 1-2 key achievements
- Connect your experience to the role

[Call to Action - 1-2 sentences]
- Request an interview or meeting
- Ask for a discussion about the position

Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.

Yours sincerely,

[Your Full Name]
Email: [Your Email]
Phone: [Your Phone Number]
LinkedIn: [Your LinkedIn URL]

CRITICAL FORMATTING RULES:
1. Each section on its own line with blank lines between
2. "Yours sincerely," on its own line with a blank line after
3. Name on its own line after "Yours sincerely,"
4. Email on its own line
5. Phone on its own line
6. LinkedIn on its own line
7. Use "Dear Hiring Manager," (not "Dear hr")
8. Maximum 80-100 words total
9. Professional left alignment
10. Subject line should be professional and attention-grabbing
"""

    # Set output based on choice
    if choice == "1":
        output_type = cover_letter_prompt
    elif choice == "2":
        output_type = cold_email_prompt
    elif choice == "3":
        output_type = f"""
Generate TWO SEPARATE PROFESSIONAL DOCUMENTS:

DOCUMENT 1 - COVER LETTER:
{cover_letter_prompt}

DOCUMENT 2 - COLD EMAIL:
{cold_email_prompt}

IMPORTANT: Generate both documents separately. The Cover Letter should be detailed with 4-5 paragraphs. The Cold Email should be concise with 80-100 words.
"""

    # Create prompt
    prompt = f"""
{output_type}

RESUME CONTENT:
{resume_text}

JOB DESCRIPTION:
{job_description}

COMPANY: {company_name}
ROLE: {job_role}
RECRUITER: {recruiter_name}

YOUR DETAILS:
Name: [Extract from resume]
Email: [Extract from resume]
Phone: [Extract from resume]
LinkedIn: [Extract from resume or use your profile URL]

TONE: {tone}

CRITICAL: Ensure each contact detail (Email, Phone, LinkedIn) is on its own separate line.
"""

    # Generate response
    print("\n🔄 Generating Response...\n")

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.5,
            messages=[
                {
                    "role": "system",
                    "content": """
You are a professional career advisor and expert at creating job application documents.

CRITICAL FORMATTING RULES - YOU MUST FOLLOW THESE EXACTLY:

1. COVER LETTER STRUCTURE:
   [Date]
   
   [Recipient Name]
   [Recipient Title]
   [Company Name]
   [Company Address]
   
   Dear Hiring Manager,
   
   [Paragraph 1]
   
   [Paragraph 2]
   
   [Paragraph 3]
   
   [Paragraph 4]
   
   [Paragraph 5]
   
   Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.
   
   Yours sincerely,
   
   [Your Full Name]
   Email: [Your Email]
   Phone: [Your Phone Number]
   LinkedIn: [Your LinkedIn URL]

2. COLD EMAIL STRUCTURE:
   Subject: [Subject Line]
   
   Dear Hiring Manager,
   
   [Opening - 1-2 sentences]
   
   [Body - 2-3 sentences]
   
   [Call to Action - 1-2 sentences]
   
   Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.
   
   Yours sincerely,
   
   [Your Full Name]
   Email: [Your Email]
   Phone: [Your Phone Number]
   LinkedIn: [Your LinkedIn URL]

3. CONTACT DETAILS FORMAT:
   - Name on its own line
   - Email: on its own line
   - Phone: on its own line
   - LinkedIn: on its own line
   - Each must be SEPARATE lines

4. OTHER RULES:
   - Use "Dear Hiring Manager," (not "Dear hr")
   - "Yours sincerely," on its own line with a blank line after
   - Professional left alignment
   - Proper spacing between all sections
   - No titles like "COVER LETTER" or "COLD EMAIL"
"""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        generated_content = response.choices[0].message.content
        
        # Clean up the content
        generated_content = re.sub(r'Dear\s+hr[,\s]*', 'Dear Hiring Manager,\n\n', generated_content, flags=re.IGNORECASE)
        generated_content = re.sub(r'Email:\s*Email:', 'Email:', generated_content, flags=re.IGNORECASE)
        generated_content = re.sub(r'Phone:\s*Phone:', 'Phone:', generated_content, flags=re.IGNORECASE)
        generated_content = re.sub(r'LinkedIn:\s*LinkedIn:', 'LinkedIn:', generated_content, flags=re.IGNORECASE)
        generated_content = re.sub(r'www\.\s*LinkedIn:', 'LinkedIn:', generated_content, flags=re.IGNORECASE)
        generated_content = re.sub(r'Yours sYours sincerely', 'Yours sincerely', generated_content, flags=re.IGNORECASE)
        
        # Fix contact details formatting
        generated_content = fix_contact_details(generated_content)
        
        print("=" * 70)
        print("AI GENERATED OUTPUT")
        print("=" * 70)
        print()
        print(generated_content)
        print()
        print("=" * 70)
        
        # Ask if user wants to download
        print("\n💾 Do you want to download this content?")
        download_choice = get_case_insensitive_input(
            "Enter 'yes' to download or 'no' to skip (yes/no): ",
            valid_options=['yes', 'no']
        )
        
        if download_choice == 'yes':
            print("\n📥 Downloading documents...\n")
            saved_files, output_dir = save_content(generated_content, choice)
            
            if saved_files:
                print("\n" + "=" * 70)
                print("✅ DOCUMENTS SAVED SUCCESSFULLY!")
                print("=" * 70)
                print(f"📁 Location: {output_dir}")
                print("\n📄 Files created:")
                for doc_type, files in saved_files.items():
                    print(f"\n  {doc_type.replace('_', ' ').upper()}:")
                    for format_type, file_path in files.items():
                        print(f"    - {os.path.basename(file_path)} ({format_type.upper()})")
                print("\n" + "=" * 70)
                
                # Ask if user wants to open the PDF files
                print("\n📂 Do you want to open the PDF files?")
                open_choice = get_case_insensitive_input(
                    "Enter 'yes' to open PDFs or 'no' to skip (yes/no): ",
                    valid_options=['yes', 'no']
                )
                
                if open_choice == 'yes':
                    print("\n📂 Opening PDF files...")
                    opened_files = open_pdf_files(saved_files)
                    if opened_files:
                        print(f"\n✅ Opened {len(opened_files)} PDF file(s)")
                    else:
                        print("\n⚠️ Could not open PDF files")
            else:
                print("❌ Failed to save documents.")
        
        # Ask if user wants to generate again
        print("\n🔄 Do you want to generate another document?")
        again_choice = get_case_insensitive_input(
            "Enter 'yes' to continue or 'no' to exit (yes/no): ",
            valid_options=['yes', 'no']
        )
        
        if again_choice == 'yes':
            print("\n" + "="*70 + "\n")
            main()
        else:
            print("\n" + "=" * 70)
            print("Thank you for using AI Application Generator!")
            print("=" * 70)

    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()