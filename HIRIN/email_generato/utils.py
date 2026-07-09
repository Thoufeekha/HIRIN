# email_generator/utils.py
import fitz
import re
import os
import subprocess
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
import platform
from django.conf import settings

# Load environment variables from Email folder (matching your structure)
env_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'Email', '.env')
load_dotenv(dotenv_path=env_path)

# Get API key
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not found in .env file")

# Initialize Groq client
client = Groq(api_key=GROQ_API_KEY)

def get_downloads_folder():
    """Get the user's Downloads folder path"""
    if platform.system() == 'Windows':
        return os.path.join(os.environ['USERPROFILE'], 'Downloads')
    elif platform.system() == 'Darwin':
        return os.path.join(os.path.expanduser('~'), 'Downloads')
    else:
        return os.path.join(os.path.expanduser('~'), 'Downloads')

def open_file(file_path):
    """Open file with default application"""
    try:
        if platform.system() == 'Windows':
            os.startfile(file_path)
        elif platform.system() == 'Darwin':
            subprocess.run(['open', file_path])
        else:
            subprocess.run(['xdg-open', file_path])
        return True
    except Exception as e:
        print(f"⚠️ Could not open file: {e}")
        return False

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        print("Error Reading PDF:", e)
        return None

def extract_jd_details(jd_text):
    """Extract company, role, and recruiter from job description"""
    company = None
    role = None
    recruiter = None

    company_match = re.search(r"Company\s*:\s*(.*)", jd_text, re.IGNORECASE)
    if company_match:
        company = company_match.group(1).strip()

    role_match = re.search(r"Job Title\s*:\s*(.*)", jd_text, re.IGNORECASE)
    if role_match:
        role = role_match.group(1).strip()

    recruiter_match = re.search(r"(Hiring Manager|HR|Recruitment Team|Hiring Team)", jd_text, re.IGNORECASE)
    if recruiter_match:
        recruiter = recruiter_match.group(1)

    return company, role, recruiter

def fix_contact_details(text):
    """Fix contact details to ensure each is on its own line"""
    lines = text.split('\n')
    new_lines = []
    contact_found = False
    contact_line = ""
    
    for line in lines:
        if "Yours sincerely" in line or "Yours sincerely," in line:
            new_lines.append(line)
            contact_found = True
            continue
        
        if contact_found:
            if "Email:" in line or "Phone:" in line or "LinkedIn:" in line or "@" in line or "+" in line:
                contact_line += " " + line.strip()
            else:
                new_lines.append(line)
        else:
            new_lines.append(line)
    
    if contact_line:
        contact_line = contact_line.strip()
        
        email_match = re.search(r'Email:\s*([^\s]+)', contact_line, re.IGNORECASE)
        email = email_match.group(0) if email_match else ""
        
        phone_match = re.search(r'Phone:\s*([^\s]+(?:\s*[^\s]+)*)', contact_line, re.IGNORECASE)
        phone = phone_match.group(0) if phone_match else ""
        
        linkedin_match = re.search(r'LinkedIn:\s*([^\s]+)', contact_line, re.IGNORECASE)
        linkedin = linkedin_match.group(0) if linkedin_match else ""
        
        if not email and '@' in contact_line:
            email_match = re.search(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})', contact_line)
            if email_match:
                email = f"Email: {email_match.group(1)}"
        
        if not phone and '+' in contact_line:
            phone_match = re.search(r'(\+\d{1,3}\s*\d{10})', contact_line)
            if phone_match:
                phone = f"Phone: {phone_match.group(1)}"
        
        if not linkedin and 'linkedin.com' in contact_line:
            linkedin_match = re.search(r'(linkedin\.com/[^\s]+)', contact_line)
            if linkedin_match:
                linkedin = f"LinkedIn: {linkedin_match.group(1)}"
        
        contact_section = []
        if email:
            contact_section.append(email)
        if phone:
            contact_section.append(phone)
        if linkedin:
            contact_section.append(linkedin)
        
        if contact_section:
            new_lines = [line for line in new_lines if not (('Email:' in line or 'Phone:' in line or 'LinkedIn:' in line or '@' in line or '+' in line) and not 'Yours sincerely' in line)]
            for contact in contact_section:
                new_lines.append(contact)
    
    final_text = '\n'.join(new_lines)
    final_text = re.sub(r'www\.linkedin\.com', 'LinkedIn: www.linkedin.com', final_text, flags=re.IGNORECASE)
    final_text = re.sub(r'linkedin\.com', 'LinkedIn: linkedin.com', final_text, flags=re.IGNORECASE)
    
    return final_text

def format_cover_letter(text):
    """Format Cover Letter with proper professional structure"""
    text = re.sub(r'COVER LETTER', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Cover Letter', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Dear\s+hr[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Dear\s+Hiring\s+Manager[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sincerely,?\s*([A-Z][a-z]+ [A-Z][a-z]+)', r'Yours sincerely,\n\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sYours sincerely', 'Yours sincerely', text, flags=re.IGNORECASE)
    text = fix_contact_details(text)
    text = re.sub(r'---\s*SEPARATOR\s*---', '', text, flags=re.IGNORECASE)
    text = re.sub(r'SEPARATOR', '', text, flags=re.IGNORECASE)
    text = re.sub(r'-{3,}', '', text)
    text = re.sub(r'\. ([A-Z])', r'. \n\n\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'www\.\s*LinkedIn:', 'LinkedIn:', text, flags=re.IGNORECASE)
    text = re.sub(r' +', ' ', text)
    return text.strip()

def format_cold_email(text):
    """Format Cold Email with proper professional structure"""
    text = re.sub(r'COLD EMAIL', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Cold Email', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Subject:\s*', 'Subject: ', text, flags=re.IGNORECASE)
    if not text.startswith('Subject:'):
        text = 'Subject: Application for Position\n\n' + text
    text = re.sub(r'Dear\s+hr[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Dear\s+Hiring\s+Manager[,\s]*', 'Dear Hiring Manager,\n\n', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sincerely,?\s*([A-Z][a-z]+ [A-Z][a-z]+)', r'Yours sincerely,\n\n\1', text, flags=re.IGNORECASE)
    text = re.sub(r'Yours sYours sincerely', 'Yours sincerely', text, flags=re.IGNORECASE)
    text = fix_contact_details(text)
    text = re.sub(r'---\s*SEPARATOR\s*---', '', text, flags=re.IGNORECASE)
    text = re.sub(r'SEPARATOR', '', text, flags=re.IGNORECASE)
    text = re.sub(r'-{3,}', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'www\.\s*LinkedIn:', 'LinkedIn:', text, flags=re.IGNORECASE)
    text = re.sub(r' +', ' ', text)
    return text.strip()

def save_as_docx(content, filename, content_type="Cover Letter"):
    """Save content as DOCX file with professional formatting"""
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = 72
        section.bottom_margin = 72
        section.left_margin = 72
        section.right_margin = 72
        
        if "Cover Letter" in content_type:
            content = format_cover_letter(content)
        else:
            content = format_cold_email(content)
        
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if re.match(r'^[A-Za-z]+\s+\d{1,2},?\s+\d{4}$', para.strip()):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif para.strip().startswith('Dear'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif para.strip().startswith('Subject:'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif para.strip().startswith('Yours sincerely'):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif any(keyword in para.strip() for keyword in ['Email:', 'Phone:', 'LinkedIn:']):
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = 36
                    p.paragraph_format.space_after = 12
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.save(filename)
        print(f"✅ {content_type} saved as: {filename}")
        return True
    except Exception as e:
        print(f"❌ Error saving {content_type} as DOCX: {e}")
        return False

def save_as_pdf(content, filename, content_type="Cover Letter"):
    """Save content as PDF file with professional formatting"""
    try:
        doc = SimpleDocTemplate(
            filename, 
            pagesize=letter,
            leftMargin=72,
            rightMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        styles = getSampleStyleSheet()
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            leading=14,
            firstLineIndent=36
        )
        
        left_style = ParagraphStyle(
            'LeftAlign',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=6,
            leading=14
        )
        
        if "Cover Letter" in content_type:
            content = format_cover_letter(content)
        else:
            content = format_cold_email(content)
        
        story = []
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if re.match(r'^[A-Za-z]+\s+\d{1,2},?\s+\d{4}$', para.strip()):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                elif para.strip().startswith('Dear'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                elif para.strip().startswith('Subject:'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                elif para.strip().startswith('Yours sincerely'):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.1*inch))
                elif any(keyword in para.strip() for keyword in ['Email:', 'Phone:', 'LinkedIn:']):
                    story.append(Paragraph(para.strip(), left_style))
                    story.append(Spacer(1, 0.05*inch))
                else:
                    story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 0.05*inch))
        
        doc.build(story)
        print(f"✅ {content_type} saved as: {filename}")
        return True
    except Exception as e:
        print(f"❌ Error saving {content_type} as PDF: {e}")
        return False

def save_cover_letter(content, output_dir):
    """Save cover letter in both DOCX and PDF formats"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    docx_filename = os.path.join(output_dir, f"cover_letter_{timestamp}.docx")
    pdf_filename = os.path.join(output_dir, f"cover_letter_{timestamp}.pdf")
    
    docx_success = save_as_docx(content, docx_filename, "Cover Letter")
    pdf_success = save_as_pdf(content, pdf_filename, "Cover Letter")
    
    if docx_success and pdf_success:
        return {
            'docx': docx_filename,
            'pdf': pdf_filename
        }
    return None

def save_cold_email(content, output_dir):
    """Save cold email in both DOCX and PDF formats"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    docx_filename = os.path.join(output_dir, f"cold_email_{timestamp}.docx")
    pdf_filename = os.path.join(output_dir, f"cold_email_{timestamp}.pdf")
    
    docx_success = save_as_docx(content, docx_filename, "Cold Email")
    pdf_success = save_as_pdf(content, pdf_filename, "Cold Email")
    
    if docx_success and pdf_success:
        return {
            'docx': docx_filename,
            'pdf': pdf_filename
        }
    return None

def extract_cover_letter_and_email(content):
    """Extract Cover Letter and Cold Email from combined content"""
    cover_letter = None
    cold_email = None
    
    if "SEPARATOR" in content:
        parts = re.split(r'---?\s*SEPARATOR\s*---?', content, flags=re.IGNORECASE)
        if len(parts) >= 2:
            cover_letter = parts[0].strip()
            cold_email = parts[1].strip()
            return cover_letter, cold_email
    
    if "Subject:" in content:
        parts = content.split("Subject:", 1)
        if len(parts) >= 2:
            cover_letter = parts[0].strip()
            cold_email = "Subject:" + parts[1].strip()
            return cover_letter, cold_email
    
    dear_indices = []
    lines = content.split('\n')
    for i, line in enumerate(lines):
        if "Dear" in line:
            dear_indices.append(i)
    
    if len(dear_indices) >= 2:
        cover_letter = '\n'.join(lines[:dear_indices[1]]).strip()
        cold_email = '\n'.join(lines[dear_indices[1]:]).strip()
        return cover_letter, cold_email
    
    return content, None

def save_content(content, choice, output_dir):
    """Save generated content to files based on choice"""
    saved_files = {}
    
    if choice == "cover_letter":
        result = save_cover_letter(content, output_dir)
        if result:
            saved_files['cover_letter'] = result
    
    elif choice == "cold_email":
        result = save_cold_email(content, output_dir)
        if result:
            saved_files['cold_email'] = result
    
    elif choice == "both":
        cover_letter, cold_email = extract_cover_letter_and_email(content)
        
        if cover_letter:
            result = save_cover_letter(cover_letter, output_dir)
            if result:
                saved_files['cover_letter'] = result
        
        if cold_email:
            result = save_cold_email(cold_email, output_dir)
            if result:
                saved_files['cold_email'] = result
        
        if not saved_files:
            result = save_cover_letter(content, output_dir)
            if result:
                saved_files['cover_letter'] = result
            result = save_cold_email(content, output_dir)
            if result:
                saved_files['cold_email'] = result
    
    return saved_files

def generate_documents(resume_text, job_description, company_name, job_role, recruiter_name, tone, document_type):
    """Generate cover letter and/or cold email using AI - Matches your original code"""
    
    result = {'cover_letter': None, 'cold_email': None}
    
    # System prompt matching your original code
    system_prompt = """
You are a professional career advisor and expert at creating job application documents.

CRITICAL FORMATTING RULES - YOU MUST FOLLOW THESE EXACTLY:

1. COVER LETTER STRUCTURE:
   [Date]
   
   [Recipient Name]
   [Recipient Title]
   [Company Name]
   [Company Address]
   
   Dear Hiring Manager,
   
   [Paragraph 1]
   
   [Paragraph 2]
   
   [Paragraph 3]
   
   [Paragraph 4]
   
   [Paragraph 5]
   
   Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.
   
   Yours sincerely,
   
   [Your Full Name]
   Email: [Your Email]
   Phone: [Your Phone Number]
   LinkedIn: [Your LinkedIn URL]

2. COLD EMAIL STRUCTURE:
   Subject: [Subject Line]
   
   Dear Hiring Manager,
   
   [Opening - 1-2 sentences]
   
   [Body - 2-3 sentences]
   
   [Call to Action - 1-2 sentences]
   
   Thank you for your time and consideration. I look forward to the opportunity to discuss my application further.
   
   Yours sincerely,
   
   [Your Full Name]
   Email: [Your Email]
   Phone: [Your Phone Number]
   LinkedIn: [Your LinkedIn URL]

3. CONTACT DETAILS FORMAT:
   - Name on its own line
   - Email: on its own line
   - Phone: on its own line
   - LinkedIn: on its own line
   - Each must be SEPARATE lines

4. OTHER RULES:
   - Use "Dear Hiring Manager," (not "Dear hr")
   - "Yours sincerely," on its own line with a blank line after
   - Professional left alignment
   - Proper spacing between all sections
   - No titles like "COVER LETTER" or "COLD EMAIL"
"""
    
    try:
        # Generate Cover Letter
        if document_type in ['cover_letter', 'both']:
            print("Generating Cover Letter...")
            
            cover_prompt = f"""
Generate a PROFESSIONAL COVER LETTER with proper business letter format.

RESUME CONTENT:
{resume_text[:1500]}

JOB DESCRIPTION:
{job_description[:1000]}

COMPANY: {company_name}
ROLE: {job_role}
RECRUITER: {recruiter_name}
TONE: {tone}

Generate a cover letter with:
1. Current date
2. Dear Hiring Manager,
3. 4-5 professional paragraphs
4. Yours sincerely,
5. Contact details (Email, Phone, LinkedIn on separate lines)

Use a {tone.lower()} tone.
"""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                temperature=0.5,
                max_tokens=1500,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": cover_prompt}
                ]
            )
            
            result['cover_letter'] = fix_contact_details(response.choices[0].message.content)
            print("✅ Cover Letter generated")
        
        # Generate Cold Email
        if document_type in ['cold_email', 'both']:
            print("Generating Cold Email...")
            
            email_prompt = f"""
Generate a PROFESSIONAL COLD EMAIL with proper business email format.

RESUME CONTENT:
{resume_text[:1500]}

JOB DESCRIPTION:
{job_description[:1000]}

COMPANY: {company_name}
ROLE: {job_role}
RECRUITER: {recruiter_name}
TONE: {tone}

Generate a cold email with:
1. Subject line
2. Dear Hiring Manager,
3. 3-4 short paragraphs (80-100 words total)
4. Call to action
5. Yours sincerely,
6. Contact details on separate lines

Use a {tone.lower()} tone.
"""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                temperature=0.5,
                max_tokens=800,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": email_prompt}
                ]
            )
            
            result['cold_email'] = fix_contact_details(response.choices[0].message.content)
            print("✅ Cold Email generated")
        
        return result
        
    except Exception as e:
        print(f"❌ Error generating documents: {e}")
        return None